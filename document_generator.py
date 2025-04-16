import io
from docx import Document
from docx.shared import Pt

def generate_docx_with_comments(text, suggestions):
    doc = Document()
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    for suggestion in suggestions:
        comment = f"[{suggestion['rule']}] {suggestion['message']}"
        doc.add_paragraph(comment, style='Intense Quote')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
